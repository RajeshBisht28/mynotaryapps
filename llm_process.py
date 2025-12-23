
"""
# @Leaflet Technology : version 1.1.0
# Usig LLM, Fetch text via json request.
"""


import logging
import fitz
import configparser
from langchain.chains import LLMChain
#from langchain.chat_models import AzureChatOpenAI
from langchain_community.chat_models import AzureChatOpenAI
from langchain.prompts import PromptTemplate
from docx import Document
import json
import re
import os


template_json = """
You are an intelligent document parser. Extract the following fields from the document content accurately and return them in a structured JSON format. Only include the fields listed below. If a field is not present, return null.

Fields to extract:
- Contract_ID
- CrossRefID
- Legal Entity
- FileType
- Customer
- CustomerTypeId
- ProjectOffice
- SalespersonType
- CategoryTypeCode
- AcknowledgementTypeid
- Job Number
- Dollars
- Received Date
- By
- Notify Date
- Closed Date
- ReferredFromId
- AcknowledgementMethodId

Document:
{document_content}

JSON Output:
"""

class LeafletLLM_Ver10:
    def __init__(self, logger, filepath, prompt_template_json, entity_json_str, caller="WEB-API"):
        self.file_path = filepath
        self.logger = logger
        self.caller = caller
        self.prompt_template_json = prompt_template_json
        self.entity_json_str = entity_json_str        
        self.llm = self._initialize_model()
        
    
    def _setlogger(self):
        logging.basicConfig(
        filename="vehicle_chrono.log",
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        filemode='a',
        level=logging.DEBUG
    )
        logger = logging.getLogger('VehicleEntityExt_log')
        return logger

    def _create_filepath(self, filename):
        current_file_path = os.path.abspath(__file__)
        current_directory = os.path.dirname(current_file_path)
        full_path = os.path.join(current_directory, filename)
        return full_path

    
    def _initialize_model(self):
        """Initializes the AzureChatOpenAI model."""
        logger = self.logger
        
        config = self._load_config('azure_openai_config.properties')
        try:
            deployment_name_prop = config.get('azure.openai', 'deployement_name')
            openai_api_version_prop = config.get('azure.openai', 'api_version')
            OPENAI_API_KEY_prop = config.get('azure.openai', 'api_key')
            OPENAI_API_BASE_prop = config.get('azure.openai', 'api_base')
            openai_api_type_prop = config.get('azure.openai', 'api_type')

            model = AzureChatOpenAI(
                openai_api_version=openai_api_version_prop,
                openai_api_key=OPENAI_API_KEY_prop,
                openai_api_base=OPENAI_API_BASE_prop,
                openai_api_type=openai_api_type_prop,
                deployment_name=deployment_name_prop,
                temperature=0.3
            )
            return model
        except configparser.NoOptionError as e:
            logger.error(f"Configuration error in azureConfig.properties: {e}")
            raise
        except Exception as e:
            logger.error(f"Error initializing Azure OpenAI model: {e}")
            raise

    def _load_config(self, file_name):
        """Loads configuration from a specified file."""
        logger = self.logger
        file_path = self._create_filepath(file_name)
        config = configparser.ConfigParser()
        if not config.read(file_path):
            logger.error(f"Configuration file not found or empty: {file_path}")
            raise FileNotFoundError(f"Configuration file not found: {file_path}")
        return config

    def doc_processing_text(self):
        logger = self.logger
        _, extension = os.path.splitext(self.file_path)
        
        if(extension.strip().upper() == ".DOCX"):
            return self.read_docx_content()
            
        if(extension.strip().upper() == ".PDF"):
            return self.read_pdf_content()
        
        logger.error(f"document format not supported: {self.file_path}")
        return "document format not supported"
        
    def read_docx_content(self):
        file_content = ""
        logger = self.logger
        
        try:
            doc = Document(self.file_path)
            full_text = []

            for element in doc.element.body:
                if element.tag.endswith('p'): 
                    for p in doc.paragraphs:
                        if p._element == element:
                            if p.text.strip():
                                full_text.append(p.text.strip())
                            break

                elif element.tag.endswith('tbl'): 
                    for t in doc.tables:
                        if t._element == element:
                            for row in t.rows:
                                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                                if row_text:
                                    full_text.append(row_text)
                            break

            file_content =  "\n".join(full_text)
        except Exception as e:
            logger.debug(f"Exception read_docx_content {e} ", exc_info=True)
        finally:
            return file_content
        
    def read_pdf_content(self):
        logger = self.logger
        page_text_list = []
        doc = None
        try:
            doc = fitz.open(self.file_path)
            page_count = doc.page_count
            for page_number in range(page_count):
                pg_text = self.get_text_by_page_number(doc, page_number)
                page_text_list.append(pg_text)
        except Exception as e:
            logger.debug(f"Exception read_pdf_content {e} ", exc_info=True)
        finally:
            doc.close()
            return page_text_list
        
    def get_text_by_page_number(self, doc, page_number):
        """Extracts text content from a specific page of the PDF."""
        text = ''
        #sample_text = "out best client hyundai car model 4X with year 2025 held Contract ID as 12345690 in india  may be close on 10 JUL 2027 M-900 as Model No ruunning since 2 ears and we earned is good as 9000$ till date only."
        #return sample_text
        logger = self.logger
        try:
            
            # Check if page_number is valid
            if page_number > doc.page_count:
                logger.warning(f"Invalid page number: {page_number}. Document has {doc.page_count} pages.")
                return ""
            page = doc[page_number]
            text = page.get_text()      
        except FileNotFoundError:
            logger.error(f"PDF file not found: {self.file_path}")
        except fitz.FileDataError:
            logger.error(f"Could not open PDF file, possibly corrupted: {self.file_path}")
        except Exception as e:
            logger.error(f"Error extracting text from page {page_number} of {self.file_path}: {e}", exc_info=True)
        return text

    def _llm_response(self, document_content_value, entity_json_str, template_string):
        """Generates a response from the LLM based on document content."""
 
        logger = self.logger
        try:
            prompt_template_json = PromptTemplate(
                template=template_string,
                input_variables=["extract_field_tag","document_content"]
            )
            
            llm_chain = LLMChain(llm=self.llm, prompt=prompt_template_json)
            
            response = llm_chain.run(
            extract_field_tag=entity_json_str,
            document_content=document_content_value
            )
            
            #logger.debug(f"response: {response}")
            clean_json = re.sub(r"^```json\s*|\s*```$", "", response.strip())
            #logger.debug(f"Raw LLM response: {clean_json}")
            
            try:
                json_response = json.loads(clean_json)
                return json_response
            except json.JSONDecodeError:
                logger.warning(f"LLM response was not a valid JSON: {clean_json}")
                return {"error": "LLM response not valid JSON"}

        except Exception as e:
            logger.error(f"Error during LLM chain execution: {e}", exc_info=True)
            return {"error": str(e)}
        finally:
            pass

    
    def request_entity(self):
        """Requests vehicle entity extraction for a given page number."""
        logger = self.logger
        logger.debug(f"request_entity started - caller: {self.caller}")
        all_results = []
        try:
            page_text_list = self.doc_processing_text()
            
            if not page_text_list:
                logger.warning(f"Pages extracted from file is EMPTY.")
                return {"status": "No content Found"}
            #pg_number = 1
            logger.debug(f"Total pages : {len(page_text_list)}")
            #for content in page_text_list:
            all_text = ' '.join(page_text_list)
            text = ' '.join(all_text.split())
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            content = ''.join(char for char in text if char.isprintable())
            response = self._llm_response(content, self.entity_json_str, self.prompt_template_json)
            #response['pagenumber'] = str(pg_number)
            #pg_number = pg_number + 1
            #all_results.append(response)
            result_json = json.dumps(response, indent=2)
            return result_json
        except Exception as e:
            logger.error(f"Error in request_entity: {e}", exc_info=True)
            return {"error": str(e)}
        finally:
            logger.debug("request_entity Completed...") 
        

if __name__ == "__main__":
    file_path_to_use = r"F:\IISSites\ddd.pdf" 
    caller = "SHELL-COMMAND"

    try:
        objcls = LeafletLLM_Ver10(file_path_to_use, "", caller)
        # Attempt to extract entities from page 1
        extracted_data = objcls.request_entity()
        print("\nExtracted Vehicle Data:")
        print(json.dumps(extracted_data, indent=2))
    except Exception as e:
        print(f"An error occurred during execution: {e}")
        
